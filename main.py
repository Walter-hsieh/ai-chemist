# main.py
# NOTE: You will need to install libraries:
# pip install "fastapi[all]" uvicorn python-dotenv openai google-generativeai rdkit openpyxl python-docx pypdf

import os
import requests
import google.generativeai as genai
import openai
import xml.etree.ElementTree as ET
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from rdkit import Chem
from rdkit.Chem import Draw
import io
import base64
import openpyxl
from docx import Document
from docx.shared import Inches
from google.generativeai.types import HarmCategory, HarmBlockThreshold
import shutil
from typing import List
import pypdf

# Load environment variables from a .env file for local development
load_dotenv()

# --- Create a directory for local knowledge base ---
KNOWLEDGE_BASE_DIR = "knowledge_base"
os.makedirs(KNOWLEDGE_BASE_DIR, exist_ok=True)


# --- Pydantic Models for API data structures ---
# Base model to handle common fields for different providers
class BaseRequest(BaseModel):
    api_key: str
    model_name: str = Field(default=None)
    api_provider: str = Field(default="google")

class RefineRequest(BaseRequest):
    original_proposal: str
    user_feedback: str

class StructureRequest(BaseRequest):
    proposal_text: str

class FinalProposalRequest(BaseRequest):
    summary_text: str
    proposal_text: str
    smiles_string: str
    structure_image_base64: str
    molecule_name: str # <-- ADDED

# --- FastAPI App Setup ---
app = FastAPI()

# Add CORS middleware to allow the frontend to communicate with the backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Safety Settings for Google's AI model ---
safety_settings = {
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
}

# --- Helper Functions for Data Reading ---
def read_pdf(file_path: str) -> str:
    """Reads text from a PDF file."""
    try:
        reader = pypdf.PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return ""

def read_docx(file_path: str) -> str:
    """Reads text from a DOCX file."""
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
        return ""

def read_txt(file_path: str) -> str:
    """Reads text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""

# --- Helper Functions for Data Retrieval ---
def search_local_knowledge():
    """Reads all documents from the local knowledge base."""
    combined_text = ""
    for filename in os.listdir(KNOWLEDGE_BASE_DIR):
        file_path = os.path.join(KNOWLEDGE_BASE_DIR, filename)
        if filename.endswith(".pdf"):
            content = read_pdf(file_path)
        elif filename.endswith(".docx"):
            content = read_docx(file_path)
        elif filename.endswith(".txt"):
            content = read_txt(file_path)
        else:
            continue
        
        if content:
            combined_text += f"--- Document: {filename} ---\n{content}\n\n"
            
    if not combined_text:
        return []
    return [{'title': 'Local Knowledge Base', 'abstract': combined_text}]


def search_semantic_scholar(topic: str, limit: int = 5):
    """Searches Semantic Scholar for papers on a given topic."""
    api_key = os.environ.get("SEMANTIC_SCHOLAR_API_KEY")
    headers = {'x-api-key': api_key} if api_key else {}
    search_url = "https://api.semanticscholar.org/graph/v1/paper/search"
    params = {'query': topic, 'limit': limit, 'fields': 'title,abstract'}
    try:
        response = requests.get(search_url, params=params, headers=headers, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=503, detail=f"Semantic Scholar request failed: {e}")
    results = response.json()
    return [paper for paper in results.get('data', []) if paper and paper.get('abstract')]

def search_arxiv(topic: str, limit: int = 5):
    """Searches the arXiv API for papers on a given topic."""
    base_url = 'http://export.arxiv.org/api/query?'
    search_query = f'all:{topic}'
    params = {'search_query': search_query, 'start': 0, 'max_results': limit}
    try:
        response = requests.get(base_url, params=params, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=503, detail=f"arXiv request failed: {e}")
    root = ET.fromstring(response.content)
    papers = []
    namespace = '{http://www.w3.org/2005/Atom}'
    for entry in root.findall(f'{namespace}entry'):
        title = entry.find(f'{namespace}title').text.strip()
        abstract = entry.find(f'{namespace}summary').text.strip()
        papers.append({'title': title, 'abstract': abstract})
    return papers

# --- API Endpoints ---
@app.get("/")
def read_root():
    """Serves the main index.html file for the frontend."""
    return FileResponse('index.html')

@app.post("/api/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Handles file uploads and saves them to the knowledge base."""
    for file in files:
        file_path = os.path.join(KNOWLEDGE_BASE_DIR, file.filename)
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Could not save file: {file.filename}. Error: {e}")
        finally:
            file.file.close()
    return JSONResponse(content={"message": f"Successfully uploaded {len(files)} files."})


@app.get("/api/summarize")
def get_summary(topic: str, source: str, api_key: str, api_provider: str, model_name: str = None):
    """Generates the initial literature summary and research proposal."""
    # Fetch papers regardless of the provider
    papers = []
    if source == 'local':
        papers = search_local_knowledge()
    elif source == 'arxiv':
        papers = search_arxiv(topic)
    else: # Default to semantic
        papers = search_semantic_scholar(topic)

    if not papers: 
        if source == 'local':
            return {"topic": topic, "summary": "Your local knowledge base is empty. Please upload some documents first.", "proposal": ""}
        return {"topic": topic, "summary": "Could not find any relevant papers from the selected source.", "proposal": ""}
    
    abstracts_text = "\n\n---\n\n".join(f"Title: {p['title']}\nAbstract: {p['abstract']}" for p in papers)
    
    if source == 'local':
        summary_prompt = f"Summarize the key findings and themes from the following documents for a chemist. The user is interested in the topic: '{topic}'.\n\n{abstracts_text}"
    else:
        summary_prompt = f"Summarize these abstracts about '{topic}' for a chemist:\n{abstracts_text}"

    try:
        summary_text = ""
        proposal_text = ""
        if api_provider == 'openai':
            openai.api_key = api_key
            openai_model = model_name or 'gpt-3.5-turbo'
            
            summary_completion = openai.chat.completions.create(
                model=openai_model,
                messages=[{"role": "user", "content": summary_prompt}]
            )
            summary_text = summary_completion.choices[0].message.content
            
            proposal_prompt = f"Based on this summary, propose a novel research direction:\n{summary_text}"
            proposal_completion = openai.chat.completions.create(
                model=openai_model,
                messages=[{"role": "user", "content": proposal_prompt}]
            )
            proposal_text = proposal_completion.choices[0].message.content

        elif api_provider == 'google':
            genai.configure(api_key=api_key)
            google_model = model_name or 'gemini-2.5-flash'
            model = genai.GenerativeModel(google_model)
            
            summary_response = model.generate_content(summary_prompt, safety_settings=safety_settings)
            summary_text = summary_response.text
            
            proposal_prompt = f"Based on this summary, propose a novel research direction:\n{summary_text}"
            proposal_response = model.generate_content(proposal_prompt, safety_settings=safety_settings)
            proposal_text = proposal_response.text
        else:
            raise HTTPException(status_code=400, detail="Invalid API provider.")

        return {"topic": topic, "summary": summary_text, "proposal": proposal_text}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred with the AI provider: {e}")

def get_ai_response(request: BaseRequest, prompt: str) -> str:
    """Helper function to get a response from the selected AI provider."""
    try:
        if request.api_provider == 'openai':
            openai.api_key = request.api_key
            model = request.model_name or 'o3-mini-2025-01-31'
            completion = openai.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}]
            )
            return completion.choices[0].message.content
        
        elif request.api_provider == 'google':
            genai.configure(api_key=request.api_key)
            model_name = request.model_name or 'gemini-2.5-flash'
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt, safety_settings=safety_settings)
            return response.text
        
        else:
            raise HTTPException(status_code=400, detail="Invalid API provider.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred with the AI provider: {e}")


@app.post("/api/refine_proposal")
def refine_proposal(request: RefineRequest):
    """Generates a new proposal based on user feedback."""
    refine_prompt = f"A user disliked a proposal for this reason: '{request.user_feedback}'. Original proposal: '{request.original_proposal}'. Generate a new proposal."
    new_proposal = get_ai_response(request, refine_prompt)
    return {"new_proposal": new_proposal}


@app.post("/api/generate_structure")
def generate_structure(request: StructureRequest):
    """Generates a chemical structure image and name from a text proposal."""
    for _ in range(3): # Try 3 times
        # UPDATE THE PROMPT to ask for the name in a specific format
        smiles_prompt = f"Based on this proposal, generate a plausible SMILES string and its common or IUPAC name. Respond ONLY in the format:\nSMILES: [the SMILES string]\nNAME: [the chemical name]\n\nProposal:\n{request.proposal_text}"
        
        response_text = get_ai_response(request, smiles_prompt).strip()
        
        # PARSE the response to extract SMILES and NAME
        try:
            lines = response_text.split('\n')
            smiles_line = next(line for line in lines if line.startswith("SMILES:")).replace("SMILES:", "").strip()
            name_line = next(line for line in lines if line.startswith("NAME:")).replace("NAME:", "").strip()
            
            smiles_string = smiles_line.replace("`", "")
            mol = Chem.MolFromSmiles(smiles_string)

            if mol is not None:
                img = Draw.MolToImage(mol, size=(300, 300))
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                img_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                # RETURN the name along with the other data
                return {"image": img_base64, "smiles": smiles_string, "name": name_line}
        except (StopIteration, ValueError) as e:
            # This handles cases where the AI response isn't in the expected format
            print(f"Could not parse AI response: {response_text}. Error: {e}")
            continue # Try again

    raise HTTPException(status_code=500, detail="AI failed to generate a valid chemical structure and name after multiple attempts.")


@app.post("/api/generate_proposal")
def generate_final_proposal(request: FinalProposalRequest):
    """Generates the full suite of downloadable proposal documents."""
    full_proposal_prompt = f"You are a PhD chemist writing a detailed research proposal. Use the following information to construct it. Make sure to include sections like Introduction, Background, Proposed Research, Methodology, and Expected Outcomes. Be comprehensive.\n\nCONTEXT:\n- Literature Summary: {request.summary_text}\n- Core Idea: {request.proposal_text}\n- Target Molecule SMILES: {request.smiles_string}\n- Target Molecule Name: {request.molecule_name}" # <-- UPDATED
    full_proposal_text = get_ai_response(request, full_proposal_prompt)
    
    recipe_prompt = f"Based on the proposal, create a synthesis recipe... PROPOSAL: {request.proposal_text} TARGET SMILES: {request.smiles_string}"
    recipe_data_str = get_ai_response(request, recipe_prompt)
    
    # Create Excel files
    wb_recipe = openpyxl.Workbook()
    ws_recipe = wb_recipe.active
    ws_recipe.title = "Synthesis Recipe"
    ws_recipe.append(["Chemical Name", "Molar Mass (g/mol)", "Amount (mg or µL)", "Equivalents"])
    for row in recipe_data_str.strip().split('\n'):
        split_row = [item.strip() for item in row.split(',')]
        if len(split_row) == 4: ws_recipe.append(split_row)
    recipe_buffer = io.BytesIO()
    wb_recipe.save(recipe_buffer)
    recipe_base64 = base64.b64encode(recipe_buffer.getvalue()).decode('utf-8')
    
    wb_data = openpyxl.Workbook()
    ws_data = wb_data.active
    ws_data.title = "Experimental Data"
    ws_data.append(["Experiment ID", "Date", "Reactant 1 (mg)", "Reactant 2 (mg)", "Solvent (mL)", "Reaction Time (h)", "Temperature (°C)", "Yield (%)", "Notes"])
    data_buffer = io.BytesIO()
    wb_data.save(data_buffer)
    data_base64 = base64.b64encode(data_buffer.getvalue()).decode('utf-8')

    # Create Word Document
    doc = Document()
    doc.add_heading('AI-Generated Research Proposal', 0)
    
    try:
        if request.structure_image_base64:
            img_data_url = request.structure_image_base64
            header, encoded = img_data_url.split(",", 1)
            image_data = base64.b64decode(encoded)
            image_stream = io.BytesIO(image_data)
            doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(3.0))
            doc.add_paragraph(f"Figure 1: Proposed Molecular Structure - {request.molecule_name}", style='Caption') # <-- UPDATED
    except Exception as e:
        print(f"Could not add image to docx: {e}")
    
    for line in full_proposal_text.split('\n'):
        if line.startswith('### '): doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '): doc.add_heading(line.replace('## ', ''), level=2)
        elif line.strip(): doc.add_paragraph(line)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_base64 = base64.b64encode(doc_buffer.getvalue()).decode('utf-8')

    return JSONResponse(content={
        "full_proposal_text": full_proposal_text,
        "recipe_file": recipe_base64,
        "data_template_file": data_base64,
        "proposal_docx_file": doc_base64
    })